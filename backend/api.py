import os
import uuid
import requests
from typing import Union, List, Dict, Any
from dotenv import load_dotenv

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# LangChain - OpenAI
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser

# Import Schema va Prompts tu file schema.py
from schema import HRM_SCHEMA_ENHANCED, ANSWER_PROMPT, get_schema_by_role, get_sql_prompt_by_role

# ==========================================================
# 1. SETUP & CAU HINH
# ==========================================================
load_dotenv()

# BAT BUOC phai co OpenAI API Key
if not os.environ.get("OPENAI_API_KEY"):
    raise RuntimeError("Chua cau hinh OPENAI_API_KEY")

HRM_API_URL = "https://hrm.icss.com.vn/ICSS/api/execute-sql"

app = FastAPI(title="ICS HRM SQL Chatbot API", version="3.0 - OpenAI")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Tao thu muc luu file tam
EXPORT_DIR = "./static/reports"
if not os.path.exists(EXPORT_DIR):
    os.makedirs(EXPORT_DIR)

# ==========================================================
# 2. KHOI TAO LLM (OPENAI)
# ==========================================================
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0,
    max_tokens=600
)

# ==========================================================
# 3. PYDANTIC MODELS (Request / Response)
# ==========================================================
class ConversationMessage(BaseModel):
    role: str  # 'user' or 'bot'
    content: str

class ChatRequest(BaseModel):
    question: str
    user_id: Union[int, None] = None
    role: Union[str, None] = None  # 'admin', 'manager', 'employee'
    phong_ban_id: Union[int, None] = None
    conversation_history: Union[List[ConversationMessage], None] = None  # Context Memory

class LoginRequest(BaseModel):
    username: str
    password: str

class LoginResponse(BaseModel):
    success: bool
    message: str
    user: Union[Dict, None] = None

class ChatResponse(BaseModel):
    sql: Union[str, None]
    data: Union[List, Dict, Any, None]
    answer: str
    download_url: Union[str, None] = None

class BriefingRequest(BaseModel):
    user_id: int
    role: str  # 'admin', 'manager', 'employee'
    phong_ban_id: Union[int, None] = None

class BriefingResponse(BaseModel):
    greeting: str
    checkin_status: Union[Dict, None] = None
    tasks_today: Union[List, None] = None
    leave_balance: Union[Dict, None] = None
    alerts: Union[List, None] = None
    team_summary: Union[Dict, None] = None  # Cho manager
    dept_tasks_summary: Union[Dict, None] = None  # Cho manager
    dept_projects_summary: Union[Dict, None] = None  # Cho manager
    company_summary: Union[Dict, None] = None  # Cho admin

# ==========================================================
# 4. HAM TAO BAO CAO WORD
# ==========================================================
def create_word_report(data, title="BÁO CÁO HRM", filename_prefix="report", question="", summary=""):
    if not data: return None
    
    if isinstance(data, dict):
        data = [data]
    
    doc = Document()
    
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    if question:
        doc.add_heading("1. Yêu cầu truy vấn", level=1)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f'"{question}"')
        q_run.font.italic = True
        q_run.font.size = Pt(11)
        doc.add_paragraph()
    
    if summary:
        doc.add_heading("2. Tóm tắt kết quả", level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.paragraph_format.space_after = Pt(12)
        doc.add_paragraph()
    
    section_num = 3 if question and summary else (2 if question or summary else 1)
    doc.add_heading(f"{section_num}. Dữ liệu chi tiết ({len(data)} bản ghi)", level=1)
    
    headers = list(data[0].keys())
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h).upper().replace('_', ' ')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        
    for item in data:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            cell_value = item.get(h, '')
            row_cells[i].text = str(cell_value) if cell_value is not None else ''
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("-" * 50)
    footer_run.font.color.rgb = RGBColor(200, 200, 200)
    
    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = footer_info.add_run("Báo cáo được tạo tự động bởi ICS HRM Chatbot")
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
            
    filename = f"{filename_prefix}_{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    
    return filepath

# ==========================================================
# 5. HELPER FUNCTIONS
# ==========================================================
def validate_sql(sql: str) -> str:
    sql_clean = sql.replace("```sql", "").replace("```", "").strip()
    
    forbidden = ["insert", "update", "delete", "drop", "alter", "truncate", "grant"]
    if any(cmd in sql_clean.lower() for cmd in forbidden):
        print(f"Blocked dangerous SQL: {sql_clean}")
        return ""
    
    # FIX AUTOMATIC: Replace ho_ten = '...' with ho_ten LIKE '%...%' (Common LLM mistake)
    import re
    # Pattern: ho_ten = 'value' or nv.ho_ten = 'value' or n.ho_ten = 'value', etc.
    sql_clean = re.sub(
        r"(\w+\.)?ho_ten\s*=\s*'([^']+)'",
        lambda m: f"{m.group(1) or ''}ho_ten LIKE '%{m.group(2)}%'",
        sql_clean,
        flags=re.IGNORECASE
    )
    
    return sql_clean

def execute_sql_api(sql: str) -> Any:
    if not sql: return None

    print(f"\n[DEBUG SQL]: {sql}")

    try:
        payload = {"command": sql}
        res = requests.post(HRM_API_URL, json=payload, timeout=30)
        
        if res.status_code == 200:
            try:
                result = res.json()
                # Kiểm tra nếu server trả về lỗi
                if isinstance(result, dict) and result.get('success') == False:
                    error_msg = result.get('error', 'Unknown error')
                    print(f"[API REJECTED]: {error_msg}")
                    print(f"[PROBLEM SQL]: {sql}")
                return result
            except:
                return res.text
        else:
            print(f"API Error {res.status_code}: {res.text}")
            return f"Lỗi từ hệ thống dữ liệu: {res.text}"
    except Exception as e:
        print(f"Connection Error: {e}")
        return "Lỗi kết nối đến máy chủ dữ liệu."

# ==========================================================
# 6. DAILY BRIEFING ENDPOINT
# ==========================================================
@app.post("/briefing", response_model=BriefingResponse)
async def get_daily_briefing(req: BriefingRequest):
    """
    API lấy thông tin tóm tắt hàng ngày cho user.
    Trả về thông tin khác nhau tùy theo role.
    """
    try:
        user_id = req.user_id
        role = req.role
        dept_id = req.phong_ban_id
        
        print(f"\n[BRIEFING] User: {user_id}, Role: {role}, Dept: {dept_id}")
        
        # Lấy thông tin user
        user_sql = f"SELECT ho_ten, chuc_vu FROM nhanvien WHERE id = {user_id}"
        user_result = execute_sql_api(user_sql)
        user_name = "Bạn"
        if isinstance(user_result, dict) and user_result.get('data'):
            user_name = user_result['data'][0].get('ho_ten', 'Bạn')
        elif isinstance(user_result, list) and len(user_result) > 0:
            user_name = user_result[0].get('ho_ten', 'Bạn')
        
        # Xác định lời chào theo thời gian
        hour = datetime.now().hour
        if hour < 12:
            time_greeting = "Chào buổi sáng"
        elif hour < 18:
            time_greeting = "Chào buổi chiều"
        else:
            time_greeting = "Chào buổi tối"
        
        greeting = f"☀️ {time_greeting}, {user_name}!"
        
        # === THÔNG TIN CHUNG CHO TẤT CẢ ROLE ===
        
        # 1. Trạng thái check-in hôm nay (chỉ dành cho Employee & Manager)
        checkin_status = None
        if role != 'admin':
            checkin_sql = f"""
            SELECT check_in 
            FROM cham_cong 
            WHERE nhan_vien_id = {user_id} AND ngay = CURDATE()
            """
            checkin_result = execute_sql_api(checkin_sql)
            
            if isinstance(checkin_result, dict) and checkin_result.get('data'):
                data = checkin_result['data']
                if len(data) > 0:
                    check_in = data[0].get('check_in', '')
                    is_late = check_in and check_in >= '08:06:00'
                    checkin_status = {
                        "checked_in": bool(check_in),
                        "check_in_time": check_in,
                        "check_out_time": None,
                        "is_late": is_late,
                        "status_text": "Đi muộn" if is_late else "Đúng giờ" if check_in else "Chưa check-in"
                    }
            elif isinstance(checkin_result, list) and len(checkin_result) > 0:
                check_in = checkin_result[0].get('check_in', '')
                is_late = check_in and check_in >= '08:06:00'
                checkin_status = {
                    "checked_in": bool(check_in),
                    "check_in_time": check_in,
                    "check_out_time": None,
                    "is_late": is_late,
                    "status_text": "Đi muộn" if is_late else "Đúng giờ" if check_in else "Chưa check-in"
                }
            
            if not checkin_status:
                checkin_status = {
                    "checked_in": False,
                    "check_in_time": None,
                    "check_out_time": None,
                    "is_late": False,
                    "status_text": "Chưa check-in"
                }
        else:
            # Admin: không cần check-in cá nhân
            checkin_status = None
        
        # 2. Công việc cần làm hôm nay
        tasks_sql = f"""
        SELECT cv.ten_cong_viec, cv.han_hoan_thanh, cv.muc_do_uu_tien, cv.trang_thai
        FROM cong_viec cv
        JOIN cong_viec_nguoi_nhan cvnn ON cv.id = cvnn.cong_viec_id
        WHERE cvnn.nhan_vien_id = {user_id} 
        AND cv.trang_thai != 'Đã hoàn thành'
        ORDER BY cv.muc_do_uu_tien DESC, cv.han_hoan_thanh ASC
        LIMIT 5
        """
        tasks_result = execute_sql_api(tasks_sql)
        tasks_today = []
        
        if isinstance(tasks_result, dict) and tasks_result.get('data'):
            tasks_today = tasks_result['data']
        elif isinstance(tasks_result, list):
            tasks_today = tasks_result
        
        # 3. Số ngày phép còn lại
        leave_sql = f"""
        SELECT tong_ngay_phep, ngay_phep_da_dung, ngay_phep_con_lai
        FROM ngay_phep_nam
        WHERE nhan_vien_id = {user_id} AND nam = YEAR(CURDATE())
        """
        leave_result = execute_sql_api(leave_sql)
        leave_balance = None
        
        if isinstance(leave_result, dict) and leave_result.get('data'):
            data = leave_result['data']
            if len(data) > 0:
                leave_balance = data[0]
        elif isinstance(leave_result, list) and len(leave_result) > 0:
            leave_balance = leave_result[0]
        
        
        alerts = []
        team_summary = None
        dept_tasks_summary = None
        dept_projects_summary = None
        company_summary = None
        
        # === THÔNG TIN CHO MANAGER ===
        if role == 'manager' and dept_id:
            # Tình hình phòng ban
            team_checkin_sql = f"""
            SELECT 
                (SELECT COUNT(*) FROM nhanvien WHERE phong_ban_id = {dept_id}) as total,
                (SELECT COUNT(DISTINCT c.nhan_vien_id) 
                 FROM cham_cong c 
                 JOIN nhanvien nv ON c.nhan_vien_id = nv.id 
                 WHERE nv.phong_ban_id = {dept_id} AND c.ngay = CURDATE()) as checked_in,
                (SELECT COUNT(DISTINCT dnp.nhan_vien_id) 
                 FROM don_nghi_phep dnp 
                 JOIN nhanvien nv ON dnp.nhan_vien_id = nv.id 
                 WHERE nv.phong_ban_id = {dept_id} 
                 AND CURDATE() BETWEEN dnp.tu_ngay AND dnp.den_ngay 
                 AND dnp.trang_thai = 'da_duyet') as on_leave
            """
            team_result = execute_sql_api(team_checkin_sql)
            
            if isinstance(team_result, dict) and team_result.get('data'):
                data = team_result['data'][0] if team_result['data'] else {}
                team_summary = {
                    "total_employees": data.get('total', 0),
                    "checked_in": data.get('checked_in', 0),
                    "on_leave": data.get('on_leave', 0),
                    "not_checked_in": data.get('total', 0) - data.get('checked_in', 0) - data.get('on_leave', 0)
                }
            elif isinstance(team_result, list) and len(team_result) > 0:
                data = team_result[0]
                team_summary = {
                    "total_employees": data.get('total', 0),
                    "checked_in": data.get('checked_in', 0),
                    "on_leave": data.get('on_leave', 0),
                    "not_checked_in": data.get('total', 0) - data.get('checked_in', 0) - data.get('on_leave', 0)
                }
            
            # Công việc phòng ban
            dept_tasks_sql = f"""
            SELECT 
                COUNT(DISTINCT cv.id) as total_tasks,
                COUNT(DISTINCT CASE WHEN cv.trang_thai = 'Đã hoàn thành' THEN cv.id END) as completed_tasks,
                COUNT(DISTINCT CASE WHEN cv.trang_thai != 'Đã hoàn thành' AND cv.han_hoan_thanh < CURDATE() THEN cv.id END) as overdue_tasks
            FROM cong_viec cv
            JOIN cong_viec_nguoi_nhan cvnn ON cv.id = cvnn.cong_viec_id
            JOIN nhanvien nv ON cvnn.nhan_vien_id = nv.id
            WHERE nv.phong_ban_id = {dept_id}
            """
            dept_tasks_result = execute_sql_api(dept_tasks_sql)
            
            if isinstance(dept_tasks_result, dict) and dept_tasks_result.get('data'):
                data = dept_tasks_result['data'][0] if dept_tasks_result['data'] else {}
                dept_tasks_summary = {
                    "total_tasks": data.get('total_tasks', 0),
                    "completed_tasks": data.get('completed_tasks', 0),
                    "overdue_tasks": data.get('overdue_tasks', 0)
                }
            elif isinstance(dept_tasks_result, list) and len(dept_tasks_result) > 0:
                data = dept_tasks_result[0]
                dept_tasks_summary = {
                    "total_tasks": data.get('total_tasks', 0),
                    "completed_tasks": data.get('completed_tasks', 0),
                    "overdue_tasks": data.get('overdue_tasks', 0)
                }
            
            # Dự án phòng ban
            dept_name_sql = f"SELECT ten_phong FROM phong_ban WHERE id = {dept_id}"
            dept_name_result = execute_sql_api(dept_name_sql)
            dept_name = ""
            if isinstance(dept_name_result, dict) and dept_name_result.get('data'):
                dept_name = dept_name_result['data'][0].get('ten_phong', '')
            elif isinstance(dept_name_result, list) and len(dept_name_result) > 0:
                dept_name = dept_name_result[0].get('ten_phong', '')
            
            # Cải thiện: Thêm thông tin Leader và tiến độ dự án (Luật 25)
            dept_projects_sql = f"""
            SELECT 
                COUNT(DISTINCT d.id) as total_projects,
                COUNT(DISTINCT CASE WHEN d.ngay_ket_thuc < CURDATE() AND d.trang_thai_duan NOT IN ('Đã hoàn thành', 'Kết thúc', 'Tạm ngưng') THEN d.id END) as overdue_projects,
                STRING_AGG(DISTINCT CASE WHEN d.ngay_ket_thuc < CURDATE() AND d.trang_thai_duan NOT IN ('Đã hoàn thành', 'Kết thúc', 'Tạm ngưng')
                    THEN d.ten_du_an + ' (Leader: ' + ISNULL(nv.ho_ten, 'N/A') + ', Progress: ' + CAST(ISNULL(CAST(ROUND(AVG(td.phan_tram), 0) AS INT), 0) AS VARCHAR) + '%)'
                    ELSE NULL END, '; ') as overdue_projects_details
            FROM du_an d
            LEFT JOIN nhanvien nv ON d.lead_id = nv.id
            LEFT JOIN cong_viec cv ON d.id = cv.du_an_id
            LEFT JOIN cong_viec_tien_do td ON cv.id = td.cong_viec_id 
                AND td.thoi_gian_cap_nhat = (SELECT MAX(thoi_gian_cap_nhat) FROM cong_viec_tien_do WHERE cong_viec_id = cv.id)
            WHERE d.phong_ban LIKE '%{dept_name}%' 
                AND d.trang_thai_duan NOT IN ('Đã hoàn thành', 'Kết thúc')
            GROUP BY d.id, d.ten_du_an, nv.ho_ten
            """
            dept_projects_result = execute_sql_api(dept_projects_sql)
            
            if isinstance(dept_projects_result, dict) and dept_projects_result.get('data'):
                data = dept_projects_result['data'][0] if dept_projects_result['data'] else {}
                dept_projects_summary = {
                    "total_projects": data.get('total_projects', 0),
                    "overdue_projects": data.get('overdue_projects', 0),
                    "overdue_projects_details": data.get('overdue_projects_details', '')
                }
            elif isinstance(dept_projects_result, list) and len(dept_projects_result) > 0:
                data = dept_projects_result[0]
                dept_projects_summary = {
                    "total_projects": data.get('total_projects', 0),
                    "overdue_projects": data.get('overdue_projects', 0),
                    "overdue_projects_details": data.get('overdue_projects_details', '')
                }
            
            # Alerts
            if dept_tasks_summary and dept_tasks_summary.get('overdue_tasks', 0) > 0:
                alerts.append({
                    "type": "warning",
                    "message": f"Có {dept_tasks_summary['overdue_tasks']} công việc đang trễ hạn trong phòng"
                })
            
            if dept_projects_summary and dept_projects_summary.get('overdue_projects', 0) > 0:
                alerts.append({
                    "type": "warning",
                    "message": f"Có {dept_projects_summary['overdue_projects']} dự án đang trễ hạn trong phòng"
                })
            
            # Nhân viên chưa check-in
            if team_summary and team_summary.get('not_checked_in', 0) > 0:
                alerts.append({
                    "type": "info", 
                    "message": f"{team_summary['not_checked_in']} nhân viên chưa check-in hôm nay"
                })
        
        # === THÔNG TIN CHO ADMIN ===
        if role == 'admin':
            company_sql = """
            SELECT 
                (SELECT COUNT(*) FROM nhanvien WHERE trang_thai_lam_viec LIKE '%Đang%' OR trang_thai_lam_viec IS NULL) as total_employees,
                (SELECT COUNT(DISTINCT nhan_vien_id) FROM cham_cong WHERE DATE(ngay) = CURDATE()) as checked_in_today,
                (SELECT COUNT(*) FROM du_an WHERE trang_thai_duan LIKE '%Đang%' OR trang_thai_duan LIKE '%thực hiện%') as active_projects,
                (SELECT COUNT(*) FROM cong_viec WHERE trang_thai != 'Đã hoàn thành' AND han_hoan_thanh < CURDATE()) as overdue_tasks,
                (SELECT COUNT(*) FROM du_an WHERE ngay_ket_thuc < CURDATE() AND trang_thai_duan NOT IN ('Đã hoàn thành', 'Tạm ngưng')) as overdue_projects
            """
            company_result = execute_sql_api(company_sql)
            
            print(f"[BRIEFING ADMIN] Company result type: {type(company_result)}")
            print(f"[BRIEFING ADMIN] Company result: {company_result}")
            
            if isinstance(company_result, dict) and company_result.get('data'):
                data = company_result['data'][0] if company_result['data'] else {}
                company_summary = {
                    "total_employees": data.get('total_employees', 0) or 0,
                    "checked_in_today": data.get('checked_in_today', 0) or 0,
                    "active_projects": data.get('active_projects', 0) or 0,
                    "overdue_tasks": data.get('overdue_tasks', 0) or 0,
                    "overdue_projects": data.get('overdue_projects', 0) or 0
                }
            elif isinstance(company_result, list) and len(company_result) > 0:
                data = company_result[0]
                company_summary = {
                    "total_employees": data.get('total_employees', 0) or 0,
                    "checked_in_today": data.get('checked_in_today', 0) or 0,
                    "active_projects": data.get('active_projects', 0) or 0,
                    "overdue_tasks": data.get('overdue_tasks', 0) or 0,
                    "overdue_projects": data.get('overdue_projects', 0) or 0
                }
            else:
                # Fallback - query từng thứ riêng
                try:
                    total_sql = "SELECT COUNT(*) as cnt FROM nhanvien WHERE trang_thai_lam_viec LIKE '%Đang%' OR trang_thai_lam_viec IS NULL"
                    total_res = execute_sql_api(total_sql)
                    total_emp = (total_res[0].get('cnt', 0) if isinstance(total_res, list) else 
                                total_res.get('data', [{}])[0].get('cnt', 0)) if total_res else 0
                    
                    checkin_sql = "SELECT COUNT(DISTINCT nhan_vien_id) as cnt FROM cham_cong WHERE DATE(ngay) = CURDATE()"
                    checkin_res = execute_sql_api(checkin_sql)
                    checked_in = (checkin_res[0].get('cnt', 0) if isinstance(checkin_res, list) else 
                                 checkin_res.get('data', [{}])[0].get('cnt', 0)) if checkin_res else 0
                    
                    company_summary = {
                        "total_employees": total_emp or 0,
                        "checked_in_today": checked_in or 0,
                        "active_projects": 0,
                        "overdue_tasks": 0,
                        "overdue_projects": 0
                    }
                except:
                    company_summary = {
                        "total_employees": 0,
                        "checked_in_today": 0,
                        "active_projects": 0,
                        "overdue_tasks": 0,
                        "overdue_projects": 0
                    }
            
            print(f"[BRIEFING ADMIN] Final company_summary: {company_summary}")
            
            if company_summary and company_summary.get('overdue_tasks', 0) > 0:
                alerts.append({
                    "type": "warning",
                    "message": f"Có {company_summary['overdue_tasks']} công việc đang trễ hạn trong công ty"
                })
        
        return BriefingResponse(
            greeting=greeting,
            checkin_status=checkin_status,
            tasks_today=tasks_today[:5] if tasks_today else [],
            leave_balance=leave_balance,
            alerts=alerts,
            team_summary=team_summary,
            dept_tasks_summary=dept_tasks_summary,
            dept_projects_summary=dept_projects_summary,
            company_summary=company_summary
        )
        
    except Exception as e:
        print(f"[BRIEFING ERROR]: {str(e)}")
        import traceback
        traceback.print_exc()
        return BriefingResponse(
            greeting="Xin chào!",
            checkin_status=None,
            tasks_today=[],
            leave_balance=None,
            alerts=[{"type": "error", "message": "Không thể tải dữ liệu briefing"}],
            team_summary=None,
            company_summary=None
        )

# ==========================================================
# 6.1. TUẦN 2 - ACTION BOT APIs
# ==========================================================

# --- Pydantic Models for Action Bot ---
class LeaveRequestCreate(BaseModel):
    nhanvien_id: int
    tu_ngay: str  # YYYY-MM-DD
    den_ngay: str
    ly_do: str

class LeaveApproveRequest(BaseModel):
    request_id: int
    admin_id: int
    approved: bool

class TaskAssignRequest(BaseModel):
    ten_cong_viec: str
    mo_ta: str = ""
    du_an_id: Union[int, None] = None
    nguoi_nhan_ids: List[int]
    nguoi_giao_id: int
    han_hoan_thanh: str  # YYYY-MM-DD
    muc_do_uu_tien: str = "Trung bình"

# --- Leave Request Endpoint ---
@app.post("/leave-request")
async def create_leave_request(req: LeaveRequestCreate):
    """
    Tạo đơn xin nghỉ phép mới.
    Chỉ dành cho Employee và Manager.
    """
    try:
        print(f"\n[LEAVE REQUEST] NhanVien: {req.nhanvien_id}")
        print(f"[LEAVE REQUEST] Từ: {req.tu_ngay} -> Đến: {req.den_ngay}")
        print(f"[LEAVE REQUEST] Lý do: {req.ly_do}")
        
        # Tạo SQL insert
        sql = f"""
        INSERT INTO don_nghi_phep (nhanvien_id, tu_ngay, den_ngay, ly_do, trang_thai, ngay_tao)
        VALUES ({req.nhanvien_id}, '{req.tu_ngay}', '{req.den_ngay}', N'{req.ly_do}', N'Chờ duyệt', NOW())
        """
        
        result = execute_sql_api(sql)
        print(f"[LEAVE REQUEST] Result: {result}")
        
        # Demo mode fallback
        if isinstance(result, str) and "Lỗi" in result:
            return {
                "success": True,
                "message": "Đơn nghỉ phép đã được gửi thành công (Demo)",
                "demo_mode": True
            }
        
        return {
            "success": True,
            "message": "Đơn nghỉ phép đã được gửi thành công",
            "demo_mode": False
        }
        
    except Exception as e:
        print(f"[LEAVE REQUEST ERROR]: {e}")
        return {
            "success": True,  # Return success for demo mode
            "message": "Đơn nghỉ phép đã được gửi (Demo mode)",
            "demo_mode": True
        }

# ==========================================================
# 7. ADMIN ANALYTICS DASHBOARD ENDPOINT
# ==========================================================

class AnalyticsResponse(BaseModel):
    stats: Dict[str, Any]
    task_completion_rate: float
    top_employees: List[Dict[str, Any]]
    employee_workload: List[Dict[str, Any]]
    project_health: List[Dict[str, Any]]
    department_stats: List[Dict[str, Any]]
    hourlyData: List[Dict[str, Any]]
    timestamp: str

@app.get("/admin/analytics", response_model=AnalyticsResponse)
async def get_admin_analytics():
    """
    API lấy dữ liệu thống kê cho Admin Dashboard
    Trả về: totalEmployees, checkedInToday, totalTasks, completedTasks, overdueTasks, activeProjects
    """
    try:
        # 1. Basic Stats (Reusing some queries)
        stats_sql = """
        SELECT
            (SELECT COUNT(*) FROM nhanvien WHERE trang_thai_lam_viec = 'Đang làm') as total_employees,
            (SELECT COUNT(DISTINCT nhan_vien_id) FROM cham_cong WHERE DATE(ngay) = CURDATE() AND check_in IS NOT NULL) as checked_in_today,
            (SELECT COUNT(*) FROM cong_viec) as total_tasks,
            (SELECT COUNT(*) FROM cong_viec WHERE trang_thai = 'Đã hoàn thành') as completed_tasks,
            (SELECT COUNT(*) FROM cong_viec WHERE trang_thai != 'Đã hoàn thành' AND han_hoan_thanh < CURDATE()) as overdue_tasks,
            (SELECT COUNT(*) FROM du_an WHERE trang_thai_duan = 'Đang thực hiện') as active_projects
        """
        stats_result = execute_sql_api(stats_sql)
        stats = stats_result['data'][0] if isinstance(stats_result, dict) and stats_result.get('data') else {}

        # 2. Tỉ Lệ Hoàn Thành Task (%)
        total_tasks = stats.get('total_tasks', 0)
        completed_tasks = stats.get('completed_tasks', 0)
        task_completion_rate = round((completed_tasks / total_tasks * 100), 1) if total_tasks > 0 else 0

        # 3. Top 5 Nhân Viên Xuất Sắc (hoàn thành nhiều task nhất)
        top_employees_sql = """
        SELECT
            nv.ho_ten,
            pb.ten_phong,
            COUNT(cv.id) as completed_tasks
        FROM nhanvien nv
        JOIN cong_viec_nguoi_nhan cvnn ON nv.id = cvnn.nhan_vien_id
        JOIN cong_viec cv ON cvnn.cong_viec_id = cv.id
        LEFT JOIN phong_ban pb ON nv.phong_ban_id = pb.id
        WHERE cv.trang_thai = 'Đã hoàn thành'
        GROUP BY nv.ho_ten, pb.ten_phong
        ORDER BY completed_tasks DESC
        LIMIT 5;
        """
        top_employees_result = execute_sql_api(top_employees_sql)
        top_employees = top_employees_result['data'] if isinstance(top_employees_result, dict) and top_employees_result.get('data') else []

        # 4. Workload Per Employee (số task đang active)
        employee_workload_sql = """
        SELECT
            nv.ho_ten,
            pb.ten_phong,
            COUNT(cv.id) as active_tasks
        FROM nhanvien nv
        LEFT JOIN cong_viec_nguoi_nhan cvnn ON nv.id = cvnn.nhan_vien_id
        LEFT JOIN cong_viec cv ON cvnn.cong_viec_id = cv.id AND cv.trang_thai NOT IN ('Đã hoàn thành', 'Tạm ngưng')
        WHERE nv.trang_thai_lam_viec = 'Đang làm'
        GROUP BY nv.ho_ten, pb.ten_phong
        ORDER BY active_tasks DESC;
        """
        employee_workload_result = execute_sql_api(employee_workload_sql)
        employee_workload = employee_workload_result['data'] if isinstance(employee_workload_result, dict) and employee_workload_result.get('data') else []

        # 5. Projects Health Status
        project_health_sql = """
        SELECT
            ten_du_an,
            trang_thai_duan,
            ngay_bat_dau,
            ngay_ket_thuc,
            CASE
                WHEN trang_thai_duan = 'Đã hoàn thành' THEN 'Completed'
                WHEN ngay_ket_thuc < CURDATE() AND trang_thai_duan NOT IN ('Đã hoàn thành', 'Tạm ngưng') THEN 'Overdue'
                WHEN DATEDIFF(ngay_ket_thuc, CURDATE()) <= 7 AND trang_thai_duan NOT IN ('Đã hoàn thành', 'Tạm ngưng') THEN 'At Risk'
                ELSE 'On Track'
            END as health_status
        FROM du_an
        WHERE trang_thai_duan != 'Tạm ngưng';
        """
        project_health_result = execute_sql_api(project_health_sql)
        project_health = project_health_result['data'] if isinstance(project_health_result, dict) and project_health_result.get('data') else []

        # 6. Department Statistics
        department_stats_sql = """
        SELECT
            pb.ten_phong,
            COUNT(DISTINCT nv.id) as number_of_employees,
            COUNT(DISTINCT cv.id) as total_tasks,
            COUNT(DISTINCT CASE WHEN cv.trang_thai = 'Đã hoàn thành' THEN cv.id END) as completed_tasks
        FROM phong_ban pb
        LEFT JOIN nhanvien nv ON pb.id = nv.phong_ban_id
        LEFT JOIN cong_viec_nguoi_nhan cvnn ON nv.id = cvnn.nhan_vien_id
        LEFT JOIN cong_viec cv ON cvnn.cong_viec_id = cv.id
        GROUP BY pb.ten_phong
        ORDER BY number_of_employees DESC;
        """
        department_stats_result = execute_sql_api(department_stats_sql)
        department_stats = department_stats_result['data'] if isinstance(department_stats_result, dict) and department_stats_result.get('data') else []

        # 7. Dữ liệu chấm công theo giờ (giữ lại từ code cũ)
        hourly_sql = """
        SELECT
            DATE_FORMAT(check_in, '%H:00') as hour,
            COUNT(id) as count
        FROM cham_cong
        WHERE ngay = CURDATE() AND check_in IS NOT NULL
        GROUP BY hour
        ORDER BY hour;
        """
        hourly_result = execute_sql_api(hourly_sql)
        hourly_data = hourly_result['data'] if isinstance(hourly_result, dict) and hourly_result.get('data') else []


        return {
            "stats": stats,
            "task_completion_rate": task_completion_rate,
            "top_employees": top_employees,
            "employee_workload": employee_workload,
            "project_health": project_health,
            "department_stats": department_stats,
            "hourlyData": hourly_data,
            "timestamp": datetime.now().isoformat()
        }

    except Exception as e:
        print(f"[ADMIN ANALYTICS ERROR]: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================================
# 7B. MANAGER ANALYTICS DASHBOARD ENDPOINT
# ==========================================================
@app.get("/manager/analytics")
async def get_manager_analytics(user_id: int, dept_id: int):
    """
    API lấy dữ liệu thống kê cho Manager Dashboard (chỉ dữ liệu phòng ban)
    Trả về: totalEmployees, checkedInToday, totalTasks, completedTasks, overdueTasks, activeProjects
    """
    try:
        # 1. Tổng số nhân viên trong phòng
        total_emp_sql = f"SELECT COUNT(*) as cnt FROM nhanvien WHERE phong_ban_id = {dept_id} AND trang_thai_lam_viec = 'Đang làm'"
        total_emp_result = execute_sql_api(total_emp_sql)
        total_employees = 0
        if isinstance(total_emp_result, dict) and total_emp_result.get('data'):
            total_employees = total_emp_result['data'][0].get('cnt', 0)
        elif isinstance(total_emp_result, list) and len(total_emp_result) > 0:
            total_employees = total_emp_result[0].get('cnt', 0)
        
        # 2. Check-in hôm nay (chỉ nhân viên trong phòng)
        checkin_sql = f"""
        SELECT COUNT(DISTINCT c.nhan_vien_id) as cnt 
        FROM cham_cong c
        JOIN nhanvien nv ON c.nhan_vien_id = nv.id
        WHERE DATE(c.ngay) = CURDATE() AND c.check_in IS NOT NULL AND nv.phong_ban_id = {dept_id}
        """
        checkin_result = execute_sql_api(checkin_sql)
        checked_in_today = 0
        
        try:
            if isinstance(checkin_result, dict) and checkin_result.get('data'):
                data_list = checkin_result['data']
                if isinstance(data_list, list) and len(data_list) > 0:
                    checked_in_today = int(data_list[0].get('cnt', 0) or 0)
            elif isinstance(checkin_result, list) and len(checkin_result) > 0:
                checked_in_today = int(checkin_result[0].get('cnt', 0) or 0)
        except Exception as e:
            print(f"[MANAGER ANALYTICS ERROR] Parsing checkin: {e}")
            checked_in_today = 0
        
        # 3. Công việc (chỉ nhân viên trong phòng)
        task_sql = f"""
        SELECT 
            COUNT(DISTINCT CASE WHEN cv.trang_thai = 'Đã hoàn thành' THEN cv.id END) as completed,
            COUNT(DISTINCT CASE WHEN cv.trang_thai != 'Đã hoàn thành' AND cv.han_hoan_thanh < CURDATE() THEN cv.id END) as overdue,
            COUNT(DISTINCT cv.id) as total
        FROM cong_viec cv
        JOIN cong_viec_nguoi_nhan cvnn ON cv.id = cvnn.cong_viec_id
        JOIN nhanvien nv ON cvnn.nhan_vien_id = nv.id
        WHERE nv.phong_ban_id = {dept_id}
        """
        task_result = execute_sql_api(task_sql)
        total_tasks = 0
        completed_tasks = 0
        overdue_tasks = 0
        
        if isinstance(task_result, dict) and task_result.get('data'):
            data = task_result['data'][0]
            total_tasks = data.get('total', 0)
            completed_tasks = data.get('completed', 0)
            overdue_tasks = data.get('overdue', 0)
        elif isinstance(task_result, list) and len(task_result) > 0:
            data = task_result[0]
            total_tasks = data.get('total', 0)
            completed_tasks = data.get('completed', 0)
            overdue_tasks = data.get('overdue', 0)
        
        # 4. Dự án (chỉ dự án của phòng)
        dept_name_sql = f"SELECT ten_phong FROM phong_ban WHERE id = {dept_id}"
        dept_name_result = execute_sql_api(dept_name_sql)
        dept_name = ""
        if isinstance(dept_name_result, dict) and dept_name_result.get('data'):
            dept_name = dept_name_result['data'][0].get('ten_phong', '')
        elif isinstance(dept_name_result, list) and len(dept_name_result) > 0:
            dept_name = dept_name_result[0].get('ten_phong', '')
        
        project_sql = f"""
        SELECT COUNT(*) as cnt 
        FROM du_an 
        WHERE trang_thai_duan = 'Đang thực hiện' AND phong_ban LIKE '%{dept_name}%'
        """
        project_result = execute_sql_api(project_sql)
        active_projects = 0
        
        try:
            if isinstance(project_result, dict) and project_result.get('data'):
                data_list = project_result['data']
                if isinstance(data_list, list) and len(data_list) > 0:
                    active_projects = int(data_list[0].get('cnt', 0) or 0)
            elif isinstance(project_result, list) and len(project_result) > 0:
                active_projects = int(project_result[0].get('cnt', 0) or 0)
        except Exception as e:
            print(f"[MANAGER ANALYTICS ERROR] Parsing project: {e}")
            active_projects = 0
        
        # 5. Tính % Check-in và Hoàn thành
        checked_in_percent = round((checked_in_today / total_employees * 100) if total_employees > 0 else 0)
        completed_percent = round((completed_tasks / total_tasks * 100) if total_tasks > 0 else 0)
        
        return {
            "stats": {
                "totalEmployees": total_employees,
                "checkedInToday": checked_in_today,
                "checkedInPercent": checked_in_percent,
                "totalTasks": total_tasks,
                "completedTasks": completed_tasks,
                "overdueTasks": overdue_tasks,
                "activeProjects": active_projects
            },
            "timestamp": datetime.now().isoformat()
        }
    
    except Exception as e:
        print(f"[MANAGER ANALYTICS ERROR]: {e}")
        return {
            "error": str(e),
            "stats": {
                "totalEmployees": 0,
                "checkedInToday": 0,
                "checkedInPercent": 0,
                "totalTasks": 0,
                "completedTasks": 0,
                "overdueTasks": 0,
                "activeProjects": 0
            }
        }

# --- Get Leave Requests Endpoint (for Admin) ---
@app.get("/leave-requests")
async def get_leave_requests(status: str = "pending"):
    """
    Lấy danh sách đơn nghỉ phép.
    Chỉ dành cho Admin.
    """
    try:
        if status == "pending":
            condition = "dnp.trang_thai = N'Chờ duyệt'"
        else:
            condition = "1=1"
        
        sql = f"""
        SELECT 
            dnp.id,
            dnp.nhanvien_id as nhan_vien_id,
            nv.ho_ten,
            pb.ten_phong as phong_ban,
            dnp.tu_ngay,
            dnp.den_ngay,
            DATEDIFF(dnp.den_ngay, dnp.tu_ngay) + 1 as so_ngay,
            dnp.ly_do,
            dnp.trang_thai,
            dnp.ngay_tao
        FROM don_nghi_phep dnp
        JOIN nhanvien nv ON dnp.nhanvien_id = nv.id
        LEFT JOIN phong_ban pb ON nv.phong_ban_id = pb.id
        WHERE {condition}
        ORDER BY dnp.ngay_tao DESC
        """
        
        result = execute_sql_api(sql)
        
        requests = []
        if isinstance(result, dict) and result.get('data'):
            requests = result['data']
        elif isinstance(result, list):
            requests = result
        
        return {
            "success": True,
            "requests": requests
        }
        
    except Exception as e:
        print(f"[GET LEAVE REQUESTS ERROR]: {e}")
        # Demo data
        return {
            "success": True,
            "requests": [
                {
                    "id": 1,
                    "nhan_vien_id": 3,
                    "ho_ten": "Lê Văn Cường",
                    "phong_ban": "Phòng Kỹ thuật",
                    "tu_ngay": "2026-02-05",
                    "den_ngay": "2026-02-07",
                    "so_ngay": 3,
                    "ly_do": "Về quê có việc gia đình",
                    "trang_thai": "Chờ duyệt",
                    "ngay_tao": "2026-02-01"
                },
                {
                    "id": 2,
                    "nhan_vien_id": 4,
                    "ho_ten": "Phạm Thị Dung",
                    "phong_ban": "Phòng Kinh doanh",
                    "tu_ngay": "2026-02-10",
                    "den_ngay": "2026-02-12",
                    "so_ngay": 3,
                    "ly_do": "Khám sức khỏe định kỳ",
                    "trang_thai": "Chờ duyệt",
                    "ngay_tao": "2026-02-01"
                }
            ]
        }

# --- Leave Approval Endpoint ---
@app.post("/leave-approve")
async def approve_leave_request(req: LeaveApproveRequest):
    """
    Duyệt hoặc từ chối đơn nghỉ phép.
    Chỉ dành cho Admin.
    """
    try:
        new_status = "Đã duyệt" if req.approved else "Từ chối"
        
        print(f"\n[LEAVE APPROVE] Request: {req.request_id}")
        print(f"[LEAVE APPROVE] Admin: {req.admin_id}")
        print(f"[LEAVE APPROVE] Status: {new_status}")
        
        sql = f"""
        UPDATE don_nghi_phep 
        SET trang_thai = N'{new_status}'
        WHERE id = {req.request_id}
        """
        
        result = execute_sql_api(sql)
        
        return {
            "success": True,
            "message": f"Đơn đã được {new_status.lower()}"
        }
        
    except Exception as e:
        print(f"[LEAVE APPROVE ERROR]: {e}")
        return {
            "success": True,
            "message": f"Đơn đã được xử lý (Demo mode)"
        }

# --- Get Employees for Task Assignment ---
@app.get("/employees")
async def get_employees(role: str = "admin", phong_ban_id: str = ""):
    """
    Lấy danh sách nhân viên để giao việc.
    Manager: chỉ lấy nhân viên trong phòng
    Admin: lấy tất cả
    """
    try:
        if role == "manager" and phong_ban_id:
            condition = f"nv.phong_ban_id = {phong_ban_id}"
        else:
            condition = "1=1"
        
        sql = f"""
        SELECT 
            nv.id,
            nv.ho_ten,
            pb.ten_phong as phong_ban,
            nv.chuc_vu
        FROM nhanvien nv
        LEFT JOIN phong_ban pb ON nv.phong_ban_id = pb.id
        WHERE {condition}
        AND nv.trang_thai_lam_viec = N'Đang làm việc'
        ORDER BY nv.ho_ten
        """
        
        result = execute_sql_api(sql)
        
        employees = []
        if isinstance(result, dict) and result.get('data'):
            employees = result['data']
        elif isinstance(result, list):
            employees = result
        
        return {
            "success": True,
            "employees": employees
        }
        
    except Exception as e:
        print(f"[GET EMPLOYEES ERROR]: {e}")
        # Demo data
        return {
            "success": True,
            "employees": [
                {"id": 3, "ho_ten": "Lê Văn Cường", "phong_ban": "Phòng Kỹ thuật", "chuc_vu": "Nhân viên"},
                {"id": 4, "ho_ten": "Phạm Thị Dung", "phong_ban": "Phòng Kinh doanh", "chuc_vu": "Nhân viên"},
                {"id": 6, "ho_ten": "Ngô Thị Phương", "phong_ban": "Phòng Kỹ thuật", "chuc_vu": "Nhân viên"}
            ]
        }

# --- Get Projects ---
@app.get("/projects")
async def get_projects():
    """
    Lấy danh sách dự án đang active để gán công việc.
    """
    try:
        sql = """
        SELECT id, ten_du_an, trang_thai_duan as trang_thai
        FROM du_an
        WHERE trang_thai_duan NOT LIKE N'%Hoàn thành%'
        ORDER BY ten_du_an
        """
        
        result = execute_sql_api(sql)
        
        projects = []
        if isinstance(result, dict) and result.get('data'):
            projects = result['data']
        elif isinstance(result, list):
            projects = result
        
        return {
            "success": True,
            "projects": projects
        }
        
    except Exception as e:
        print(f"[GET PROJECTS ERROR]: {e}")
        # Demo data
        return {
            "success": True,
            "projects": [
                {"id": 1, "ten_du_an": "Hệ thống quản lý nhân sự", "trang_thai": "Đang thực hiện"},
                {"id": 2, "ten_du_an": "Website công ty", "trang_thai": "Đang thực hiện"},
                {"id": 3, "ten_du_an": "App mobile", "trang_thai": "Lên kế hoạch"}
            ]
        }

# --- Task Assignment Endpoint ---
@app.post("/assign-task")
async def assign_task(req: TaskAssignRequest):
    """
    Giao công việc cho nhân viên.
    Dành cho Manager và Admin.
    """
    try:
        print(f"\n[ASSIGN TASK] Tên CV: {req.ten_cong_viec}")
        print(f"[ASSIGN TASK] Người giao: {req.nguoi_giao_id}")
        print(f"[ASSIGN TASK] Người nhận: {req.nguoi_nhan_ids}")
        print(f"[ASSIGN TASK] Hạn: {req.han_hoan_thanh}")
        
        # Insert công việc
        du_an_value = req.du_an_id if req.du_an_id else "NULL"
        
        sql_cv = f"""
        INSERT INTO cong_viec (
            ten_cong_viec, mo_ta, du_an_id, nguoi_giao_id, 
            ngay_bat_dau, han_hoan_thanh, trang_thai, muc_do_uu_tien, ngay_tao
        )
        OUTPUT INSERTED.id
        VALUES (
            N'{req.ten_cong_viec}', 
            N'{req.mo_ta}', 
            {du_an_value}, 
            {req.nguoi_giao_id},
            GETDATE(), 
            '{req.han_hoan_thanh}', 
            N'Chưa bắt đầu', 
            N'{req.muc_do_uu_tien}', 
            GETDATE()
        )
        """
        
        result_cv = execute_sql_api(sql_cv)
        
        # Lấy id công việc vừa tạo
        cv_id = None
        if isinstance(result_cv, dict) and result_cv.get('data'):
            cv_id = result_cv['data'][0].get('id')
        elif isinstance(result_cv, list) and len(result_cv) > 0:
            cv_id = result_cv[0].get('id')
        
        # Insert người nhận
        if cv_id:
            for nhan_vien_id in req.nguoi_nhan_ids:
                sql_nn = f"""
                INSERT INTO cong_viec_nguoi_nhan (cong_viec_id, nhan_vien_id)
                VALUES ({cv_id}, {nhan_vien_id})
                """
                execute_sql_api(sql_nn)
        
        return {
            "success": True,
            "message": "Công việc đã được giao thành công",
            "cong_viec_id": cv_id
        }
        
    except Exception as e:
        print(f"[ASSIGN TASK ERROR]: {e}")
        return {
            "success": True,
            "message": "Công việc đã được giao thành công (Demo mode)",
            "demo_mode": True
        }

# ==========================================================
# 7. DOWNLOAD FILE ENDPOINT
# ==========================================================
@app.get("/download/{filename}")
async def download_file(filename: str):
    if "../" in filename or "..\\" in filename:
        raise HTTPException(status_code=400, detail="Tên tệp không hợp lệ")
    
    filepath = os.path.join(EXPORT_DIR, filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

# ==========================================================
# 7. LOGIN ENDPOINT
# ==========================================================
@app.post("/login", response_model=LoginResponse)
async def login_endpoint(req: LoginRequest):
    try:
        print(f"\n{'='*50}")
        print(f"[LOGIN] Đang xử lý đăng nhập...")
        print(f"[LOGIN] Username: {req.username}")
        
        username_clean = req.username.strip()
        password_clean = req.password.strip().replace(' ', '').replace('.', '').replace('-', '')
        
        sql = f"""
        SELECT id, ho_ten, email, so_dien_thoai, chuc_vu, vai_tro, phong_ban_id 
        FROM nhanvien 
        WHERE email = '{username_clean}' 
           OR ho_ten LIKE '%{username_clean}%'
           OR LOWER(email) = LOWER('{username_clean}')
        """
        
        print(f"[LOGIN SQL]: {sql}")
        result = execute_sql_api(sql)
        
        if isinstance(result, str) and "Lỗi" in result:
            print(f"[LOGIN] ỗi kết nối DB: {result}")
            return LoginResponse(
                success=False,
                message="Không thể kết nối đến hệ thống. Vui lòng thử lại sau.",
                user=None
            )
        
        if isinstance(result, dict) and 'data' in result:
            users_data = result.get('data', [])
        elif isinstance(result, list):
            users_data = result
        else:
            users_data = []
        
        print(f"[LOGIN] Kiểu dữ liệu result: {type(result)}")
        print(f"[LOGIN] Nội dung result: {result}")
        
        if not users_data or len(users_data) == 0:
            print(f"[LOGIN] Không tìm thấy user với username: {username_clean}")
            return LoginResponse(
                success=False,
                message="Không tìm thấy tài khoản. Vui lòng kiểm tra lại email hoặc họ tên.",
                user=None
            )
        
        print(f"[LOGIN] Tìm thấy {len(users_data)} kết quả")
        
        user_found = None
        for user in users_data:
            if not isinstance(user, dict):
                print(f"[LOGIN] Bỏ qua item không phải dict: {user}")
                continue
                
            phone_raw = user.get('so_dien_thoai', '') or ''
            phone_clean = str(phone_raw).replace(' ', '').replace('.', '').replace('-', '')
            
            print(f"[LOGIN] So sánh password với SDT: '{password_clean}' vs '{phone_clean}'")
            
            if password_clean == phone_clean:
                user_found = user
                print(f"[LOGIN] Xác thực thành công cho: {user.get('ho_ten')}")
                break
        
        if not user_found:
            print(f"[LOGIN] Mật khẩu không khớp")
            return LoginResponse(
                success=False,
                message="Mật khẩu không chính xác. (ợi ý: Mật khẩu là số điện thoại của bạn)",
                user=None
            )
        
        vai_tro = user_found.get('vai_tro', '') or 'Nhân viên'
        chuc_vu = user_found.get('chuc_vu', '') or ''
        
        print(f"[LOGIN] Vai trò trong DB: '{vai_tro}'")
        print(f"[LOGIN] Chức vụ trong DB: '{chuc_vu}'")
        
        role = 'employee'
        
        # Hỗ trợ cả có dấu và không dấu
        admin_keywords = [
            'admin', 'giam doc', 'giám đốc', 'ceo', 'director', 
            'chu tich', 'chủ tịch', 'tổng giám đốc', 'pho giam doc', 'phó giám đốc'
        ]
        manager_keywords = [
            'quan ly', 'quản lý', 'manager', 
            'truong phong', 'trưởng phòng', 
            'truong nhom', 'trưởng nhóm',
            'leader', 'supervisor', 'team lead'
        ]
        
        check_text = (vai_tro + ' ' + chuc_vu).lower()
        
        for keyword in admin_keywords:
            if keyword in check_text:
                role = 'admin'
                break
        
        if role != 'admin':
            for keyword in manager_keywords:
                if keyword in check_text:
                    role = 'manager'
                    break
        
        print(f"[LOGIN] Role được gán: {role}")
        print(f"{'='*50}\n")
        
        return LoginResponse(
            success=True,
            message=f"Đăng nhập thành công! Xin chào {user_found.get('ho_ten', '')}",
            user={
                "id": user_found.get('id'),
                "ho_ten": user_found.get('ho_ten'),
                "email": user_found.get('email'),
                "chuc_vu": chuc_vu,
                "vai_tro": vai_tro,
                "role": role,
                "phong_ban_id": user_found.get('phong_ban_id')
            }
        )
        
    except Exception as e:
        print(f"[LOGIN ERROR]: {str(e)}")
        import traceback
        traceback.print_exc()
        return LoginResponse(
            success=False,
            message="Lỗi hệ thống. Vui lòng thử lại sau.",
            user=None
        )

# ==========================================================
# 8. HELPER: Kiểm tra nhân viên có thuộc phòng ban không
# ==========================================================
def check_employee_in_department(question: str, dept_id: int) -> tuple:
    """
    Kiểm tra nếu câu hỏi đề cập đến tên người cụ thể,
    xác minh người đó có thuộc phòng ban của quản lý không.
    
    Returns:
        (is_valid, message) - True nếu hợp lệ hoặc không có tên cụ thể
    """
    if not dept_id:
        return (True, None)
    
    # Dùng LLM để trích xuất tên người từ câu hỏi
    extract_prompt = f"""Phân tích câu hỏi sau và trích xuất TÊN NGƯỜI (nếu có).
Câu hỏi: "{question}"

Quy tắc:
- Nếu câu hỏi đề cập đến một người CỤ THỂ (có họ tên), trả về tên đó.
- Nếu câu hỏi hỏi chung (ai, mọi người, nhân viên nào...), trả về: NONE
- Chỉ trả về tên hoặc NONE, không giải thích.

Tên người (hoặc NONE):"""
    
    try:
        name_result = llm.invoke(extract_prompt).content.strip()
        print(f"[CHECK] Tên trích xuất: {name_result}")
        
        if name_result == "NONE" or not name_result or len(name_result) < 2:
            return (True, None)  # Không có tên cụ thể, cho phép tiếp tục
        
        # Kiểm tra người này có thuộc phòng ban không
        check_sql = f"""
        SELECT id, ho_ten, phong_ban_id 
        FROM nhanvien 
        WHERE ho_ten LIKE '%{name_result}%' 
        AND phong_ban_id = {dept_id}
        """
        print(f"[CHECK SQL]: {check_sql}")
        
        result = execute_sql_api(check_sql)
        
        # Xử lý kết quả
        if isinstance(result, dict) and 'data' in result:
            data = result.get('data', [])
        elif isinstance(result, list):
            data = result
        else:
            data = []
        
        if not data or len(data) == 0:
            # Không tìm thấy trong phòng ban
            return (False, f"Nhân viên '{name_result}' không thuộc phòng ban của bạn hoặc không tồn tại trong hệ thống.")
        
        return (True, None)  # Tìm thấy, cho phép tiếp tục
        
    except Exception as e:
        print(f"[CHECK ERROR]: {e}")
        return (True, None)  # Lỗi thì cho qua

# ==========================================================
# 9. MAIN CHAT ENDPOINT
# ==========================================================

def build_conversation_context(history: list) -> str:
    """Build conversation context string from history for LLM"""
    if not history or len(history) == 0:
        return "Không có ngữ cảnh trước đó. Đây là câu hỏi đầu tiên."
    
    # Lấy tối đa 6 tin nhắn gần nhất (3 cặp hỏi-đáp)
    recent_history = history[-6:] if len(history) > 6 else history
    
    context_parts = []
    for i, msg in enumerate(recent_history):
        role_label = "User hỏi" if msg.role == "user" else "Bot trả lời"
        # Truncate long messages
        content = msg.content[:200] + "..." if len(msg.content) > 200 else msg.content
        context_parts.append(f"{role_label}: {content}")
    
    return "\\n".join(context_parts)


@app.post("/chat", response_model=ChatResponse)
async def chat_endpoint(req: ChatRequest):
    try:
        # Chọn schema phù hợp với role của user
        role = req.role or 'employee'  # Mặc định là employee nếu không có role
        user_id = req.user_id
        dept_id = req.phong_ban_id
        
        # Lấy schema phân quyền
        user_schema = get_schema_by_role(role=role, user_id=user_id, dept_id=dept_id)
        
        # Build conversation context for Context Memory
        conversation_context = build_conversation_context(req.conversation_history or [])
        
        print(f"[CHAT] Role: {role}, User ID: {user_id}, Dept ID: {dept_id}")
        print(f"[CONTEXT] {conversation_context[:100]}...")
        
        # === KIỂM TRA QUYỀN TRUY CẬP NHÂN VIÊN (CHỈ CHO MANAGER) ===
        if role == 'manager' and dept_id:
            is_valid, error_msg = check_employee_in_department(req.question, dept_id)
            if not is_valid:
                print(f"[PERMISSION DENIED]: {error_msg}")
                return ChatResponse(
                    sql=None,
                    data=None,
                    answer=error_msg,
                    download_url=None
                )
        
        # Lấy SQL_PROMPT phù hợp với role
        sql_prompt = get_sql_prompt_by_role(role=role)
        sql_chain = sql_prompt | llm | StrOutputParser()
        raw_sql = sql_chain.invoke({
            "schema": user_schema,
            "question": req.question,
            "conversation_context": conversation_context
        })
        sql = validate_sql(raw_sql)
        
        print(f"[RAW SQL] {raw_sql[:200]}")
        print(f"[VALIDATED SQL] {sql[:200] if sql else 'EMPTY'}")

        # Kiểm tra nếu AI từ chối do không có quyền
        if "NO_PERMISSION" in sql:
            return ChatResponse(
                sql=None,
                data=None,
                answer="Xin lỗi, bạn không có quyền truy cập thông tin này.",
                download_url=None
            )

        if "NO_DATA" in sql:
            return ChatResponse(
                sql=None,
                data=None,
                answer="Xin lỗi. Tôi không có dữ liệu về vấn đề này!",
                download_url=None
            )

        if not sql:
            data_result = None
            final_answer = "Xin lỗi, tôi không thể hiểu yêu cầu này."
            download_url = None
        else:
            data_result = execute_sql_api(sql)
            print(f"[DATA RESULT] {str(data_result)[:300]}")
            print(f"[DATA TYPE] {type(data_result)}")
            print(f"[DATA IS EMPTY] {not data_result if not isinstance(data_result, str) else 'N/A'}")
            download_url = None
            
            if isinstance(data_result, str) and "Loi" in data_result:
                final_answer = f"Warning: {data_result}"
            else:
                # Extract dữ liệu thực tế từ API response
                actual_data = data_result
                if isinstance(data_result, dict) and 'data' in data_result:
                    actual_data = data_result.get('data', [])
                    print(f"[EXTRACTED DATA] {actual_data}")
                
                # Kiểm tra số lượng items
                data_count = 0
                if isinstance(actual_data, list):
                    data_count = len(actual_data)
                    print(f"[ITEM COUNT] {data_count} items")
                
                # Thêm prefix để bắt LLM nhận thức được số lượng
                data_with_count = f"[{data_count} items] {str(actual_data)}"
                
                ans_chain = ANSWER_PROMPT | llm | StrOutputParser()
                final_answer = ans_chain.invoke({
                    "question": req.question,
                    "data": data_with_count,
                    "role": role,
                    "dept_id": dept_id or "N/A"
                })
                print(f"[ANSWER] {final_answer[:200]}")
            
            q_lower = req.question.lower()
            
            if data_result and not isinstance(data_result, str):
                if "word" in q_lower or "docx" in q_lower or "van ban" in q_lower or "xuat" in q_lower or "file" in q_lower:
                    try:
                        file_path = create_word_report(
                            data=data_result, 
                            title="BÁO CÁO TRUY VẤN HRM", 
                            filename_prefix="baocao",
                            question=req.question,
                            summary=final_answer
                        )
                        if file_path:
                            filename = os.path.basename(file_path)
                            download_url = f"/download/{filename}"
                    except Exception as e:
                        print(f"Error creating word report: {e}")

        return ChatResponse(
            sql=sql,
            data=data_result,
            answer=final_answer,
            download_url=download_url
        )

    except Exception as e:
        print(f"Server Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
# ==========================================================
# 10. GET REAL USERS FOR LOGIN (DEBUG)
# ==========================================================
@app.get("/debug/users")
async def get_users_debug():
    """
    DEBUG ENDPOINT - Lấy danh sách nhân viên thực từ database
    Dùng để biết ai có thể đăng nhập và mật khẩu là gì
    """
    try:
        sql = """
        SELECT id, ho_ten, email, so_dien_thoai, chuc_vu, vai_tro, phong_ban_id 
        FROM nhanvien 
        LIMIT 20
        """
        
        result = execute_sql_api(sql)
        
        if isinstance(result, dict) and 'data' in result:
            users_data = result.get('data', [])
        elif isinstance(result, list):
            users_data = result
        else:
            users_data = []
        
        # Format dễ đọc
        formatted_users = []
        for user in users_data:
            if isinstance(user, dict):
                formatted_users.append({
                    "id": user.get('id'),
                    "ho_ten": user.get('ho_ten'),
                    "email": user.get('email'),
                    "so_dien_thoai": user.get('so_dien_thoai'),
                    "chuc_vu": user.get('chuc_vu'),
                    "vai_tro": user.get('vai_tro'),
                    "phong_ban_id": user.get('phong_ban_id'),
                    "login_hint": f"Username: {user.get('email')}, Password: {user.get('so_dien_thoai')}"
                })
        
        return {
            "status": "success",
            "total": len(formatted_users),
            "users": formatted_users,
            "note": "Dùng email hoặc ho_ten làm username, so_dien_thoai làm password"
        }
        
    except Exception as e:
        print(f"Debug Error: {e}")
        return {
            "status": "error",
            "message": str(e)
        }